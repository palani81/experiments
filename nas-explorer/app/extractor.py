"""Content and metadata extraction from various file types."""

import os
import json
import subprocess
import logging
from pathlib import Path

logger = logging.getLogger("nas_explorer.extractor")


# ─── Text Extraction ─────────────────────────────────────────────────────

def extract_text(filepath: str, mime_type: str) -> str | None:
    """Extract searchable text from a file based on its MIME type."""

    if mime_type == "application/pdf":
        return _extract_pdf(filepath)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(filepath)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _extract_xlsx(filepath)

    if mime_type and (mime_type.startswith("text/") or mime_type in (
        "application/json", "application/xml", "application/javascript",
        "application/x-yaml", "application/x-python",
    )):
        return _extract_plaintext(filepath)

    # Subtitle files
    ext = Path(filepath).suffix.lower()
    if ext in (".srt", ".vtt", ".ass", ".ssa", ".sub"):
        return _extract_plaintext(filepath)

    return None


def _extract_pdf(filepath: str) -> str | None:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:50]:  # Limit to 50 pages
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts) if text_parts else None
    except Exception as e:
        logger.debug(f"PDF extraction failed: {filepath}: {e}")
        return None


def _extract_docx(filepath: str) -> str | None:
    """Extract text from Word documents."""
    try:
        from docx import Document
        doc = Document(filepath)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts) if text_parts else None
    except Exception as e:
        logger.debug(f"DOCX extraction failed: {filepath}: {e}")
        return None


def _extract_xlsx(filepath: str) -> str | None:
    """Extract text from Excel spreadsheets."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filepath, read_only=True, data_only=True)
        text_parts = []
        for sheet_name in list(wb.sheetnames)[:10]:  # Limit sheets
            ws = wb[sheet_name]
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    text_parts.append(" ".join(cells))
                row_count += 1
                if row_count > 500:  # Limit rows per sheet
                    break
        wb.close()
        return "\n".join(text_parts) if text_parts else None
    except Exception as e:
        logger.debug(f"XLSX extraction failed: {filepath}: {e}")
        return None


def _extract_plaintext(filepath: str, max_bytes: int = 512 * 1024) -> str | None:
    """Extract text from plain text files (with size limit)."""
    try:
        size = os.path.getsize(filepath)
        read_size = min(size, max_bytes)
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return f.read(read_size)
    except Exception as e:
        logger.debug(f"Text extraction failed: {filepath}: {e}")
        return None


# ─── Metadata Extraction ─────────────────────────────────────────────────

def extract_metadata(filepath: str, mime_type: str) -> dict | None:
    """Extract rich metadata from media files."""

    if mime_type and mime_type.startswith("image/"):
        return _extract_image_meta(filepath)

    if mime_type and mime_type.startswith("video/"):
        return _extract_video_meta(filepath)

    if mime_type and mime_type.startswith("audio/"):
        return _extract_audio_meta(filepath)

    return None


def _extract_image_meta(filepath: str) -> dict | None:
    """Extract EXIF and basic image info."""
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS

        img = Image.open(filepath)
        meta = {
            "width": img.width,
            "height": img.height,
            "format": img.format,
            "mode": img.mode,
        }

        # Try EXIF
        try:
            exif_data = img._getexif()
            if exif_data:
                exif = {}
                for tag_id, value in list(exif_data.items())[:20]:
                    tag_name = TAGS.get(tag_id, str(tag_id))
                    # Only keep serializable values
                    if isinstance(value, (str, int, float)):
                        exif[tag_name] = value
                    elif isinstance(value, bytes) and len(value) < 100:
                        exif[tag_name] = value.hex()
                if exif:
                    meta["exif"] = exif
        except Exception:
            pass

        img.close()
        return meta
    except Exception as e:
        logger.debug(f"Image meta failed: {filepath}: {e}")
        return None


def _extract_video_meta(filepath: str) -> dict | None:
    """Extract video metadata using ffprobe."""
    try:
        result = subprocess.run(
            [
                "ffprobe", "-v", "error",
                "-show_format", "-show_streams",
                "-of", "json", filepath,
            ],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            return None

        data = json.loads(result.stdout)
        fmt = data.get("format", {})
        meta = {
            "duration_secs": float(fmt.get("duration", 0)),
            "bitrate": int(fmt.get("bit_rate", 0)),
            "format_name": fmt.get("format_long_name"),
        }

        for stream in data.get("streams", []):
            if stream.get("codec_type") == "video":
                meta["video_codec"] = stream.get("codec_name")
                meta["width"] = stream.get("width")
                meta["height"] = stream.get("height")
                fps = stream.get("r_frame_rate", "0/1")
                try:
                    num, den = fps.split("/")
                    meta["fps"] = round(int(num) / int(den), 2)
                except (ValueError, ZeroDivisionError):
                    pass
            elif stream.get("codec_type") == "audio":
                meta["audio_codec"] = stream.get("codec_name")
                meta["audio_channels"] = stream.get("channels")
                meta["sample_rate"] = stream.get("sample_rate")

        return meta
    except Exception as e:
        logger.debug(f"Video meta failed: {filepath}: {e}")
        return None


def _extract_audio_meta(filepath: str) -> dict | None:
    """Extract audio metadata (ID3 tags, duration, etc.)."""
    try:
        from mutagen import File as MutagenFile

        audio = MutagenFile(filepath)
        if audio is None:
            return None

        meta = {}

        # Duration and bitrate from info
        if hasattr(audio, "info"):
            if hasattr(audio.info, "length"):
                meta["duration_secs"] = round(audio.info.length, 2)
            if hasattr(audio.info, "bitrate"):
                meta["bitrate"] = audio.info.bitrate
            if hasattr(audio.info, "sample_rate"):
                meta["sample_rate"] = audio.info.sample_rate
            if hasattr(audio.info, "channels"):
                meta["channels"] = audio.info.channels

        # Common tags (try multiple naming conventions)
        tag_map = {
            "title": ["TIT2", "title", "\xa9nam", "TITLE"],
            "artist": ["TPE1", "artist", "\xa9ART", "ARTIST"],
            "album": ["TALB", "album", "\xa9alb", "ALBUM"],
            "genre": ["TCON", "genre", "\xa9gen", "GENRE"],
            "year": ["TDRC", "date", "\xa9day", "DATE"],
            "track": ["TRCK", "tracknumber", "trkn", "TRACKNUMBER"],
        }

        for key, tag_names in tag_map.items():
            for tag_name in tag_names:
                val = audio.get(tag_name)
                if val:
                    # Mutagen returns lists for some formats
                    if isinstance(val, list):
                        val = val[0]
                    meta[key] = str(val)
                    break

        return meta if meta else None
    except Exception as e:
        logger.debug(f"Audio meta failed: {filepath}: {e}")
        return None
